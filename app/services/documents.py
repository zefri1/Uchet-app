from __future__ import annotations

from collections import defaultdict
from datetime import datetime
from decimal import Decimal
import re
from uuid import uuid4
from typing import List, Optional

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from sqlalchemy import select, or_
from sqlalchemy.orm import Session, joinedload

from ..models import BillingPeriod, ChargeAllocation, GeneratedDocument, Tenant, LeasePlacement
from ..paths import GENERATED_DIR
from .calculations import quantize_money

GENERATED_DIR.mkdir(parents=True, exist_ok=True)

UTILITY_LABELS = {
    "heat": "Теплоэнергия",
    "electricity": "Электричество",
    "water": "Водоснабжение",
    "cleaning": "Уборщица",
    "rent": "Аренда",
}

MODE_LABELS = {
    "area": "По площади",
    "manual": "Вручную",
    "mixed": "Смешанный",
    "tariff": "По тарифу",
}


def _get_utility_label(allocation: ChargeAllocation) -> str:
    utype = allocation.utility_charge.utility_type if allocation.utility_charge else "rent"
    return UTILITY_LABELS.get(utype, utype)


def _safe_name(value: str) -> str:
    return "".join(char for char in value if char.isalnum() or char in (" ", "_", "-")).strip().replace(" ", "_")


def _safe_filename_part(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9А-Яа-яЁё._-]+", "_", value).strip("._-")
    return cleaned or "document"


def _period_label(period: BillingPeriod) -> str:
    if period.month_label:
        return period.month_label
    return f"{period.start_date.isoformat()}_{period.end_date.isoformat()}"


def _document_file_path(prefix: str, tenant: Optional[Tenant], period: BillingPeriod, extension: str) -> str:
    parts = [prefix, _period_label(period)]
    if tenant is not None:
        parts.insert(1, _safe_name(tenant.display_name))
    filename = "_".join(_safe_filename_part(part) for part in parts if part)
    unique_suffix = uuid4().hex[:8]
    return str(GENERATED_DIR / f"{filename}_{unique_suffix}.{extension}")


def _tenant_allocations(db: Session, billing_period_id: int, tenant_id: Optional[int] = None) -> List[ChargeAllocation]:
    query = (
        select(ChargeAllocation)
        .options(
            joinedload(ChargeAllocation.tenant),
            joinedload(ChargeAllocation.object),
            joinedload(ChargeAllocation.utility_charge),
            joinedload(ChargeAllocation.placement),
        )
        .where(ChargeAllocation.billing_period_id == billing_period_id)
    )
    if tenant_id is not None:
        query = query.where(ChargeAllocation.tenant_id == tenant_id)
    return db.scalars(query).all()



def _get_tenant_balance_info(db: Session, period: BillingPeriod, tenant_id: int) -> dict:
    from ..models import BillingPeriod, TenantPayment, ChargeAllocation, Tenant
    from sqlalchemy import select, func
    
    tenant = db.get(Tenant, tenant_id)
    if not tenant:
        return {"incoming": Decimal("0.00"), "allocated": Decimal("0.00"), "paid": Decimal("0.00"), "outgoing": Decimal("0.00")}
        
    prev_periods = db.scalars(
        select(BillingPeriod).where(BillingPeriod.start_date < period.start_date)
    ).all()
    prev_period_ids = [p.id for p in prev_periods]
    
    initial = Decimal(tenant.initial_balance)
    prev_allocs = Decimal("0.00")
    if prev_period_ids:
        prev_allocs = db.scalar(
            select(func.sum(ChargeAllocation.amount))
            .where(ChargeAllocation.tenant_id == tenant_id)
            .where(ChargeAllocation.billing_period_id.in_(prev_period_ids))
        ) or Decimal("0.00")
        
    prev_pays = Decimal("0.00")
    if prev_period_ids:
        prev_pays = db.scalar(
            select(func.sum(TenantPayment.amount))
            .where(TenantPayment.tenant_id == tenant_id)
            .where(TenantPayment.is_active == True)
            .where(TenantPayment.billing_period_id.in_(prev_period_ids))
        ) or Decimal("0.00")
        
    incoming = initial + prev_allocs - prev_pays
    
    curr_allocs = db.scalar(
        select(func.sum(ChargeAllocation.amount))
        .where(ChargeAllocation.tenant_id == tenant_id)
        .where(ChargeAllocation.billing_period_id == period.id)
    ) or Decimal("0.00")
    
    curr_pays = db.scalar(
        select(func.sum(TenantPayment.amount))
        .where(TenantPayment.tenant_id == tenant_id)
        .where(TenantPayment.is_active == True)
        .where(TenantPayment.billing_period_id == period.id)
    ) or Decimal("0.00")
    
    outgoing = incoming + curr_allocs - curr_pays
    
    return {
        "incoming": incoming,
        "allocated": curr_allocs,
        "paid": curr_pays,
        "outgoing": outgoing
    }

def generate_invoice_docx(db: Session, billing_period: BillingPeriod, tenant: Tenant) -> GeneratedDocument:
    allocations = _tenant_allocations(db, billing_period.id, tenant.id)
    if not allocations:
        raise ValueError("Для арендатора нет начислений в выбранном периоде.")

    document = Document()
    document.add_heading("Счет на возмещение коммунальных расходов", level=1)
    document.add_paragraph(f"Арендатор: {tenant.display_name} ({tenant.tenant_type})")
    document.add_paragraph(f"Период: {_period_label(billing_period)}")
    table = document.add_table(rows=1, cols=4)
    header = table.rows[0].cells
    header[0].text = "Объект"
    header[1].text = "Услуга"
    header[2].text = "Основание"
    header[3].text = "Сумма"
    total = Decimal("0.00")
    for allocation in allocations:
        row = table.add_row().cells
        row[0].text = allocation.object.name
        row[1].text = _get_utility_label(allocation)
        room_info = f" ({allocation.placement.rental_address})" if allocation.placement else ""
        row[2].text = f"Площадь {allocation.base_area} кв.м{room_info}"
        row[3].text = f"{allocation.amount:.2f}"
        total += Decimal(allocation.amount)
    bal = _get_tenant_balance_info(db, billing_period, tenant.id)
    document.add_paragraph(f"Начислено за текущий период: {quantize_money(bal['allocated']):.2f} руб.")
    if bal['incoming'] != 0:
        status_word = "Долг" if bal['incoming'] > 0 else "Переплата"
        document.add_paragraph(f"{status_word} на начало периода: {abs(quantize_money(bal['incoming'])):.2f} руб.")
    if bal['paid'] > 0:
        document.add_paragraph(f"Оплачено в текущем периоде: {quantize_money(bal['paid']):.2f} руб.")
    document.add_paragraph(f"Итого к оплате (исходящий баланс): {quantize_money(bal['outgoing']):.2f} руб.")

    file_path = _document_file_path("invoice", tenant, billing_period, "docx")
    document.save(file_path)
    generated = GeneratedDocument(
        billing_period_id=billing_period.id,
        tenant_id=tenant.id,
        document_type="invoice",
        file_path=str(file_path),
    )
    db.add(generated)
    db.commit()
    db.refresh(generated)
    return generated


def generate_act_docx(db: Session, billing_period: BillingPeriod, tenant: Tenant) -> GeneratedDocument:
    allocations = _tenant_allocations(db, billing_period.id, tenant.id)
    if not allocations:
        raise ValueError("Для арендатора нет начислений в выбранном периоде.")

    document = Document()
    document.add_heading("Акт оказанных услуг", level=1)
    document.add_paragraph(f"Арендатор: {tenant.display_name}")
    document.add_paragraph(f"Период: {_period_label(billing_period)}")
    grouped = defaultdict(Decimal)
    for allocation in allocations:
        label = _get_utility_label(allocation)
        room_info = f" ({allocation.placement.rental_address})" if allocation.placement else ""
        grouped[f"{allocation.object.name}{room_info} / {label}"] += Decimal(allocation.amount)
    for name, amount in grouped.items():
        document.add_paragraph(f"{name}: {quantize_money(amount):.2f} руб.")
    document.add_paragraph(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}")

    file_path = _document_file_path("act", tenant, billing_period, "docx")
    document.save(file_path)
    generated = GeneratedDocument(
        billing_period_id=billing_period.id,
        tenant_id=tenant.id,
        document_type="act",
        file_path=str(file_path),
    )
    db.add(generated)
    db.commit()
    db.refresh(generated)
    return generated


def generate_register_xlsx(db: Session, billing_period: BillingPeriod, tenant: Optional[Tenant] = None) -> GeneratedDocument:
    allocations = _tenant_allocations(db, billing_period.id, tenant_id=tenant.id if tenant else None)
    if not allocations:
        raise ValueError("Нет начислений для формирования реестра.")

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Реестр"
    
    # Ensure grid lines are visible
    sheet.sheet_view.showGridLines = True
    
    headers = [
        "Тип (ИП/ООО)", 
        "ФИО / Наименование", 
        "Телефон", 
        "Объект недвижимости", 
        "Общая площадь объекта, кв.м", 
        "Адрес аренды", 
        "Занимаемая площадь, кв.м", 
        "Услуга", 
        "Сумма, руб.", 
        "Режим распределения"
    ]
    sheet.append(headers)
    
    for allocation in allocations:
        rental_address = allocation.placement.rental_address if allocation.placement else ""
        sheet.append(
            [
                allocation.tenant.tenant_type,
                allocation.tenant.display_name,
                allocation.tenant.phone or "",
                allocation.object.name,
                float(allocation.object.total_area),
                rental_address,
                float(allocation.base_area),
                _get_utility_label(allocation),
                float(allocation.amount),
                MODE_LABELS.get(allocation.mode, allocation.mode),
            ]
        )

    # Style the headers
    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="1F4B99", end_color="1F4B99", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center")
    
    for col_idx in range(1, len(headers) + 1):
        cell = sheet.cell(row=1, column=col_idx)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align

    # Add total row
    num_allocations = len(allocations)
    total_row_idx = num_allocations + 2
    sheet.cell(row=total_row_idx, column=1, value="Итого").font = Font(name="Calibri", size=11, bold=True)
    
    sum_cell = sheet.cell(row=total_row_idx, column=9, value=f"=SUM(I2:I{total_row_idx-1})")
    sum_cell.font = Font(name="Calibri", size=11, bold=True)
    sum_cell.number_format = '#,##0.00'
    
    thin_border = Border(
        left=Side(style='thin', color='D9D9D9'),
        right=Side(style='thin', color='D9D9D9'),
        top=Side(style='thin', color='D9D9D9'),
        bottom=Side(style='thin', color='D9D9D9')
    )
    
    double_bottom_border = Border(
        top=Side(style='thin', color='000000'),
        bottom=Side(style='double', color='000000')
    )
    
    # Format data rows
    for r_idx in range(2, total_row_idx):
        # Format total area (col 5)
        sheet.cell(row=r_idx, column=5).number_format = '#,##0.00'
        # Format base area (col 7)
        sheet.cell(row=r_idx, column=7).number_format = '#,##0.00'
        # Format amount (col 9)
        sheet.cell(row=r_idx, column=9).number_format = '#,##0.00'
        
        for c_idx in range(1, len(headers) + 1):
            sheet.cell(row=r_idx, column=c_idx).border = thin_border
            
    # Add border to total row
    for c_idx in range(1, len(headers) + 1):
        sheet.cell(row=total_row_idx, column=c_idx).border = double_bottom_border

    # Auto-adjust column widths
    for col in sheet.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            val_str = str(cell.value or '')
            if val_str.startswith('='):  # skip formula length estimation
                val_str = "1234567.89"
            if len(val_str) > max_len:
                max_len = len(val_str)
        sheet.column_dimensions[col_letter].width = max(max_len + 3, 11)

    file_path = _document_file_path("register", tenant, billing_period, "xlsx")
    workbook.save(file_path)
    generated = GeneratedDocument(
        billing_period_id=billing_period.id,
        tenant_id=tenant.id if tenant else None,
        document_type="register",
        file_path=str(file_path),
    )
    db.add(generated)
    db.commit()
    db.refresh(generated)
    return generated


def generate_invoice_xlsx(db: Session, billing_period: BillingPeriod, tenant: Tenant) -> GeneratedDocument:
    allocations = _tenant_allocations(db, billing_period.id, tenant.id)
    if not allocations:
        raise ValueError("Для арендатора нет начислений в выбранном периоде.")

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Счет"
    sheet.sheet_view.showGridLines = True

    # Title
    sheet.cell(row=1, column=1, value="Счет на возмещение коммунальных расходов").font = Font(name="Calibri", size=14, bold=True)
    sheet.cell(row=2, column=1, value=f"Арендатор: {tenant.display_name} ({tenant.tenant_type})")
    sheet.cell(row=3, column=1, value=f"Период: {_period_label(billing_period)}")

    headers = [
        "Тип (ИП/ООО)", 
        "ФИО / Наименование", 
        "Телефон", 
        "Объект недвижимости", 
        "Общая площадь объекта, кв.м", 
        "Адрес аренды", 
        "Занимаемая площадь, кв.м", 
        "Услуга", 
        "Сумма, руб.", 
        "Режим распределения"
    ]
    sheet.append([])  # Row 4 is empty
    sheet.append(headers)  # Row 5

    for allocation in allocations:
        rental_address = allocation.placement.rental_address if allocation.placement else ""
        sheet.append(
            [
                allocation.tenant.tenant_type,
                allocation.tenant.display_name,
                allocation.tenant.phone or "",
                allocation.object.name,
                float(allocation.object.total_area),
                rental_address,
                float(allocation.base_area),
                _get_utility_label(allocation),
                float(allocation.amount),
                MODE_LABELS.get(allocation.mode, allocation.mode),
            ]
        )

    # Style headers
    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="1F4B99", end_color="1F4B99", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center")
    
    for col_idx in range(1, len(headers) + 1):
        cell = sheet.cell(row=5, column=col_idx)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align

    # Add total row
    num_allocations = len(allocations)
    total_row_idx = 5 + num_allocations + 1
    sheet.cell(row=total_row_idx, column=1, value="Итого").font = Font(name="Calibri", size=11, bold=True)
    
    sum_cell = sheet.cell(row=total_row_idx, column=9, value=f"=SUM(I6:I{total_row_idx-1})")
    sum_cell.font = Font(name="Calibri", size=11, bold=True)
    sum_cell.number_format = '#,##0.00'
    
    thin_border = Border(
        left=Side(style='thin', color='D9D9D9'),
        right=Side(style='thin', color='D9D9D9'),
        top=Side(style='thin', color='D9D9D9'),
        bottom=Side(style='thin', color='D9D9D9')
    )
    double_bottom_border = Border(
        top=Side(style='thin', color='000000'),
        bottom=Side(style='double', color='000000')
    )
    
    for r_idx in range(6, total_row_idx):
        sheet.cell(row=r_idx, column=5).number_format = '#,##0.00'
        sheet.cell(row=r_idx, column=7).number_format = '#,##0.00'
        sheet.cell(row=r_idx, column=9).number_format = '#,##0.00'
        for c_idx in range(1, len(headers) + 1):
            sheet.cell(row=r_idx, column=c_idx).border = thin_border
            
    for c_idx in range(1, len(headers) + 1):
        sheet.cell(row=total_row_idx, column=c_idx).border = double_bottom_border

    bal = _get_tenant_balance_info(db, billing_period, tenant.id)
    
    sheet.cell(row=total_row_idx+2, column=1, value="Детализация баланса:").font = Font(name="Calibri", size=11, bold=True)
    sheet.cell(row=total_row_idx+3, column=1, value="Входящий баланс (долг/переплата):")
    sheet.cell(row=total_row_idx+3, column=9, value=float(bal["incoming"])).number_format = '#,##0.00'
    
    sheet.cell(row=total_row_idx+4, column=1, value="Начислено за текущий период:")
    sheet.cell(row=total_row_idx+4, column=9, value=float(bal["allocated"])).number_format = '#,##0.00'
    
    sheet.cell(row=total_row_idx+5, column=1, value="Оплачено в текущем периоде:")
    sheet.cell(row=total_row_idx+5, column=9, value=float(bal["paid"])).number_format = '#,##0.00'
    
    sheet.cell(row=total_row_idx+6, column=1, value="Итого к оплате (исходящий баланс):").font = Font(name="Calibri", size=11, bold=True)
    final_cell = sheet.cell(row=total_row_idx+6, column=9, value=float(bal["outgoing"]))
    final_cell.font = Font(name="Calibri", size=11, bold=True)
    final_cell.number_format = '#,##0.00'

    # Auto widths
    for col in sheet.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            val_str = str(cell.value or '')
            if val_str.startswith('='):
                val_str = "1234567.89"
            if len(val_str) > max_len:
                max_len = len(val_str)
        sheet.column_dimensions[col_letter].width = max(max_len + 3, 11)

    file_path = _document_file_path("invoice", tenant, billing_period, "xlsx")
    workbook.save(file_path)
    generated = GeneratedDocument(
        billing_period_id=billing_period.id,
        tenant_id=tenant.id,
        document_type="invoice_xlsx",
        file_path=str(file_path),
    )
    db.add(generated)
    db.commit()
    db.refresh(generated)
    return generated


def generate_act_xlsx(db: Session, billing_period: BillingPeriod, tenant: Tenant) -> GeneratedDocument:
    allocations = _tenant_allocations(db, billing_period.id, tenant.id)
    if not allocations:
        raise ValueError("Для арендатора нет начислений в выбранном периоде.")

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Акт"
    sheet.sheet_view.showGridLines = True

    # Title
    sheet.cell(row=1, column=1, value="Акт оказанных услуг").font = Font(name="Calibri", size=14, bold=True)
    sheet.cell(row=2, column=1, value=f"Арендатор: {tenant.display_name}")
    sheet.cell(row=3, column=1, value=f"Период: {_period_label(billing_period)}")

    headers = [
        "Тип (ИП/ООО)", 
        "ФИО / Наименование", 
        "Телефон", 
        "Объект недвижимости", 
        "Общая площадь объекта, кв.м", 
        "Адрес аренды", 
        "Занимаемая площадь, кв.м", 
        "Услуга", 
        "Сумма, руб.", 
        "Режим распределения"
    ]
    sheet.append([])  # Row 4 is empty
    sheet.append(headers)  # Row 5

    for allocation in allocations:
        rental_address = allocation.placement.rental_address if allocation.placement else ""
        sheet.append(
            [
                allocation.tenant.tenant_type,
                allocation.tenant.display_name,
                allocation.tenant.phone or "",
                allocation.object.name,
                float(allocation.object.total_area),
                rental_address,
                float(allocation.base_area),
                _get_utility_label(allocation),
                float(allocation.amount),
                MODE_LABELS.get(allocation.mode, allocation.mode),
            ]
        )

    # Style headers
    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="1F4B99", end_color="1F4B99", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center")
    
    for col_idx in range(1, len(headers) + 1):
        cell = sheet.cell(row=5, column=col_idx)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align

    # Add total row
    num_allocations = len(allocations)
    total_row_idx = 5 + num_allocations + 1
    sheet.cell(row=total_row_idx, column=1, value="Итого").font = Font(name="Calibri", size=11, bold=True)
    
    sum_cell = sheet.cell(row=total_row_idx, column=9, value=f"=SUM(I6:I{total_row_idx-1})")
    sum_cell.font = Font(name="Calibri", size=11, bold=True)
    sum_cell.number_format = '#,##0.00'
    
    thin_border = Border(
        left=Side(style='thin', color='D9D9D9'),
        right=Side(style='thin', color='D9D9D9'),
        top=Side(style='thin', color='D9D9D9'),
        bottom=Side(style='thin', color='D9D9D9')
    )
    double_bottom_border = Border(
        top=Side(style='thin', color='000000'),
        bottom=Side(style='double', color='000000')
    )
    
    for r_idx in range(6, total_row_idx):
        sheet.cell(row=r_idx, column=5).number_format = '#,##0.00'
        sheet.cell(row=r_idx, column=7).number_format = '#,##0.00'
        sheet.cell(row=r_idx, column=9).number_format = '#,##0.00'
        for c_idx in range(1, len(headers) + 1):
            sheet.cell(row=r_idx, column=c_idx).border = thin_border
            
    for c_idx in range(1, len(headers) + 1):
        sheet.cell(row=total_row_idx, column=c_idx).border = double_bottom_border

    # Date of generation
    date_row_idx = total_row_idx + 2
    sheet.cell(row=date_row_idx, column=1, value=f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}").font = Font(name="Calibri", size=9, italic=True)

    # Auto widths
    for col in sheet.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            val_str = str(cell.value or '')
            if val_str.startswith('='):
                val_str = "1234567.89"
            if len(val_str) > max_len:
                max_len = len(val_str)
        sheet.column_dimensions[col_letter].width = max(max_len + 3, 11)

    file_path = _document_file_path("act", tenant, billing_period, "xlsx")
    workbook.save(file_path)
    generated = GeneratedDocument(
        billing_period_id=billing_period.id,
        tenant_id=tenant.id,
        document_type="act_xlsx",
        file_path=str(file_path),
    )
    db.add(generated)
    db.commit()
    db.refresh(generated)
    return generated

