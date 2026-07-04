from docx import Document
from openpyxl import load_workbook
from sqlalchemy import select
from app.db import SessionLocal
from app.models import GeneratedDocument
import os

def inspect_latest_docs():
    db = SessionLocal()
    try:
        # Find latest invoice docx
        inv_docx = db.scalars(
            select(GeneratedDocument)
            .where(GeneratedDocument.document_type == "invoice")
            .order_by(GeneratedDocument.created_at.desc())
        ).first()

        # Find latest invoice xlsx
        inv_xlsx = db.scalars(
            select(GeneratedDocument)
            .where(GeneratedDocument.document_type == "invoice_xlsx")
            .order_by(GeneratedDocument.created_at.desc())
        ).first()

        # Find latest act docx
        act_docx = db.scalars(
            select(GeneratedDocument)
            .where(GeneratedDocument.document_type == "act")
            .order_by(GeneratedDocument.created_at.desc())
        ).first()

        # Find latest act xlsx
        act_xlsx = db.scalars(
            select(GeneratedDocument)
            .where(GeneratedDocument.document_type == "act_xlsx")
            .order_by(GeneratedDocument.created_at.desc())
        ).first()

        print("=== INVOICE DOCX (Word) ===")
        if inv_docx and os.path.exists(inv_docx.file_path):
            doc = Document(inv_docx.file_path)
            for p in doc.paragraphs:
                if p.text.strip():
                    print(p.text)
            for table in doc.tables:
                print("Table in Invoice DOCX:")
                for row in table.rows:
                    print(" | ".join(cell.text.strip() for cell in row.cells))
        else:
            print("No invoice docx found or file missing.")

        print("\n=== INVOICE XLSX (Excel) ===")
        if inv_xlsx and os.path.exists(inv_xlsx.file_path):
            wb = load_workbook(inv_xlsx.file_path, data_only=True)
            sheet = wb.active
            for row in sheet.iter_rows(values_only=True):
                if any(row):
                    print(" | ".join(str(val) if val is not None else "" for val in row))
        else:
            print("No invoice xlsx found or file missing.")

        print("\n=== ACT DOCX (Word) ===")
        if act_docx and os.path.exists(act_docx.file_path):
            doc = Document(act_docx.file_path)
            for p in doc.paragraphs:
                if p.text.strip():
                    print(p.text)
        else:
            print("No act docx found or file missing.")

        print("\n=== ACT XLSX (Excel) ===")
        if act_xlsx and os.path.exists(act_xlsx.file_path):
            wb = load_workbook(act_xlsx.file_path, data_only=True)
            sheet = wb.active
            for row in sheet.iter_rows(values_only=True):
                if any(row):
                    print(" | ".join(str(val) if val is not None else "" for val in row))
        else:
            print("No act xlsx found or file missing.")

    finally:
        db.close()

if __name__ == "__main__":
    inspect_latest_docs()
